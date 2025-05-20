from docx import Document
from openai import OpenAI
import os
import json
from docx.shared import Pt, Twips
from dotenv import load_dotenv
import re


def infer_heading_level_by_font_and_indent(docx_path):
    doc = Document(docx_path)
    font_sizes = set()
    indent_levels = set()

    # 收集所有字体大小和缩进值
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size:
                font_sizes.add(run.font.size.pt)  # 获取字体大小（磅值）

        if paragraph.paragraph_format.left_indent:
            indent_levels.add(paragraph.paragraph_format.left_indent)  # 获取缩进（Twips）

    # 排序字体大小和缩进（从大到小）
    sorted_font_sizes = sorted(font_sizes, reverse=True)
    sorted_indents = sorted(indent_levels, reverse=True)

    # 推断层级
    structured_text = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # 获取当前段落的字体大小（取第一个 run 的字体）
        current_font_size = None
        for run in paragraph.runs:
            if run.font.size:
                current_font_size = run.font.size.pt
                break

        # 获取当前段落的缩进
        current_indent = paragraph.paragraph_format.left_indent

        # 推断层级（基于字体大小）
        font_level = 0
        if current_font_size and sorted_font_sizes:
            font_level = sorted_font_sizes.index(current_font_size) + 1  # 1-based

        # 推断层级（基于缩进）
        indent_level = 0
        if current_indent and sorted_indents:
            indent_level = sorted_indents.index(current_indent) + 1  # 1-based

        # 综合层级（优先字体，其次缩进）
        final_level = max(font_level, indent_level)

        structured_text.append({
            "text": text,
            "font_size": current_font_size,
            "indent": current_indent,
            "level": final_level,
        })

    return structured_text


def docx_to_markdown(docx_path):
    doc = Document(docx_path)
    markdown_lines = []

    # for paragraph in doc.paragraphs:
    #     text = paragraph.text.strip()
    #     if not text:
    #         continue
    #
    #     style = paragraph.style.name.lower()
    structured_text = infer_heading_level_by_font_and_indent(docx_path)
    max_level = max(item["level"] for item in structured_text)
    for para in structured_text:
        level = para['level']
        text = para['text'].strip()
        # 转换标题（Heading 1 → #, Heading 2 → ##）
        if level > 0 and level < max_level:
            markdown_lines.append(f"{'#' * level} {text}")
        # 普通段落
        else:
            markdown_lines.append(text)

    return markdown_lines, max_level


def docx_table_to_markdown(docx_path):
    doc = Document(docx_path)
    markdown_lines = []
    # 处理表格（转换为 Markdown 表格语法）
    for table in doc.tables:
        tb = []
        tb.append("\n| " + " | ".join(["Header"] * len(table.columns)) + " |")
        tb.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
        for row in table.rows:
            tb.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
        markdown_lines.append(tb)
    return markdown_lines


def llm_segment(client, content: str, engine_type='gpt-4o-2024-08-06'):
    system_prompt = f"""
# 指示
## 任务介绍
- **目标**：你的任务是根据语义将文档内容分割成主题连贯的单元。
具有相同主题的连续文本片段应分组到相同的分割单元中，并在主题转换时创建新的分割单元。
注意，切割后的结果需覆盖所有的文本片段，不能遗漏任何片段。
- **数据**：输入数据是一系列由“\n\n”分隔的连续文本片段。每个连续文本片段前面都有一个编号N，格式为：[EX N]。
### 输出格式
- 以**JSONL（JSON行）**格式输出分割结果，每个分割单元的字符数不应超过1000。每个词典代表一个片段，由一个或多个关于同一主题的连续文本片段组成。
每个词典应包括以下关键字：
-**segment_id**：此段的索引，从0开始。
-**start_exchange_number**：此段中第一个连续文本片段的编号。
-**end_exchange_number**：此段中最后一个连续文本片段的编号。
-**num_exchange**：一个整数，表示此段中连续文本片段的数量，计算如下：end_exchange_number- start_exchange_number+1。

预期输出的示例格式，请勿输出JSON格式分割结果意外事件的任何描述信息：
{{"segment_id": 0，"start_exchange_number": 0，"end_exchange_number": 5，"num_exchange": 6}}
{{"segment_id": 1，"start_exchange_number": 6，"end_exchange_number": 8，"num_exchange": 3}}
"""
    system_item = {"role": "system", "content": system_prompt}
    messages = [system_item]
    messages += [{"role": "user", "content": content}]

    completion = client.chat.completions.create(
        model=engine_type,
        messages=messages,
        temperature=1.0,
        max_tokens=4096,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    if completion.choices[0].message.content:
        return completion.choices[0].message.content.strip()
    else:
        print('OpenAI API Error!')
        return None


def record_results_in_txt(results: list, file_path: str, prefix='', suffix='', split=False, enum=False, mode='w'):
    new_path = ""
    if file_path.endswith('.docx') or file_path.endswith('.xlsx'):
        file_name = file_path.split('/')[-1][:-5]
        if prefix: # 只有当 prefix 非空时才创建目录和使用它作为路径
             os.makedirs(prefix, exist_ok=True) # 确保目录存在
             new_path = os.path.join(prefix, f"{file_name}{suffix}.txt")
        else: # 如果 prefix 为空，则在当前目录下创建文件
             new_path = f"{file_name}{suffix}.txt"


    with open(new_path, mode, encoding='utf-8') as f:
        for i, line in enumerate(results):
            if enum:
                f.write(f"Ex {i}: ")
            f.write(f"{line}\n")
            if split:
                f.write(f'--------------------------\n')
    print(f"{len(results)} results have been saved in {new_path}")


def find_uncovered_indices(covered_ranges, total_range_start=0, total_range_end=160):
    # Step 1: 合并重叠或连续的区间
    overlap = 0
    merged_ranges = []
    for current in sorted(covered_ranges, key=lambda x: x["start_exchange_number"]):
        if not merged_ranges:
            merged_ranges.append(current)
        else:
            last = merged_ranges[-1]
            # 如果当前区间与上一个区间重叠或连续，则合并
            if current["start_exchange_number"] <= last["end_exchange_number"] + 1:
                overlap += 1
                last["end_exchange_number"] = max(last["end_exchange_number"], current["end_exchange_number"])
            else:
                merged_ranges.append(current)

    # Step 2: 找出未被覆盖的区间
    uncovered = []
    prev_end = total_range_start - 1  # 初始化为起始点前一个数

    for interval in merged_ranges:
        start_no = interval["start_exchange_number"]
        end_no = interval["end_exchange_number"]

        # 如果当前区间的起点 > prev_end + 1，说明中间有未被覆盖的部分
        if start_no > prev_end + 1:
            uncovered.append((prev_end + 1, start_no - 1))
        prev_end = max(prev_end, end_no)

    # 检查最后一段是否覆盖到 total_range_end
    if prev_end < total_range_end:
        uncovered.append((prev_end + 1, total_range_end))

    # Step 3: 展开所有未被覆盖的索引（可选）
    print(f"overlap: {overlap}")
    print(f"uncovered: {uncovered}")
    return uncovered

def docx_process(object:str):
    prefix_path=f"docx_results_{object.split('/')[-1].split('.')[0]}"
    docx_path = object
    load_dotenv(os.path.expanduser("~/dot_env/openai.env"))
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY").split(",")[0], base_url=os.getenv("OPENAI_API_BASE"))
    engine_type = "qwen-plus"
    content_list, max_level = docx_to_markdown(docx_path)
    table_list = docx_table_to_markdown(docx_path)
    print(f"content_list: {len(content_list)}")
    record_results_in_txt(content_list, docx_path, prefix=prefix_path, suffix='2markdown')
    print(f"table_list: {len(table_list)}")
    record_results_in_txt(table_list, docx_path, prefix=prefix_path, suffix='2markdown', mode='a+')

    full_texts = []
    para = []
    prefix = [f"{'#' * l}" for l in range(1, max_level)]
    for i, c in enumerate(content_list):
        item = f"[EX {i}] {c}"
        c = c.strip()
        # 拆分段落，分段传入LLM
        para_flag = False
        for pre in prefix:
            if c.startswith(pre):
                para_flag = True
                break
        if para_flag:
            full_texts.append(para)
            para = [item]
        else:
            para.append(item)
    full_texts.append(para)
    full_texts += table_list

    coarse_chunk = []
    for p in full_texts:
        para_text = "\n\n".join(p)
        coarse_chunk.append({'len': len(para_text), 'text': para_text})
    record_results_in_txt(coarse_chunk, docx_path, prefix=prefix_path, suffix='2coarse_chunk', split=True, enum=True)
    print(f"coarse segment results: {len(coarse_chunk)}")
    
    # 粗粒度拆分后，使用LLM进行细粒度拆分
    seg_ids = []
    for p in coarse_chunk:
        seg_json = llm_segment(client, p['text'], engine_type)
        seg_id = seg_json.strip().split("\n")
        seg_ids += seg_id
    print(f"origin segment results: {len(seg_ids)}")
    # 读取LLM细粒度分块结果，得到有效的分块区间
    results = []
    for line in seg_ids:
        line = line.strip()
        if line and line[0] == '{' and line[-1] == '}':
            seg = json.loads(line)
            results.append(seg)
    print(f"final segment results: {len(results)}")
    record_results_in_txt(results, docx_path, prefix=prefix_path, suffix='2ids')
    
    # 处理分割结果
    chunks = []
    chunk_idx = 0
    for i, seg in enumerate(results):
        start = int(seg['start_exchange_number'])
        end = int(seg['end_exchange_number'])
        num_exchange = end - start + 1
        if num_exchange >= 40:
            print(f"error range 2: {start} - {end}")
            continue
        merge_list = content_list[start:end + 1]
        merge_text = "\n".join(merge_list)
        chunks.append({"segment_id": i, "num_exchange": num_exchange, "start_exchange_number": start,
                       "end_exchange_number": end,  "text": merge_text})
        chunk_idx += 1
    print(f"chunks: {len(chunks)}")
    

    # 处理未覆盖的索引
    # TODO 后续或可根据 start_exchange_number来编号，这样子方便判断上下文
    uncovered_list = find_uncovered_indices(chunks.copy(), total_range_start=0, total_range_end=len(content_list) - 1)
    uncovered_chunks = []
    for start, end in uncovered_list:
        merge_list = content_list[start:end + 1]
        merge_text = "\n".join(merge_list)
        uncovered_chunks.append({"segment_id": chunk_idx, "num_exchange": end - start + 1,
                                 "start_exchange_number": start, "end_exchange_number": end,  "text": merge_text})
        chunk_idx += 1

    # 处理表格
    # TODO 后续或可定位表格上下文，并据此重新编号
    table_chunks = []
    for tb in table_list:
        tb_txt = "\n".join(tb)
        table_chunks.append({"segment_id": chunk_idx, "num_exchange": 1, "start_exchange_number": 0,
                             "end_exchange_number": 0,  "text": tb_txt})
        chunk_idx += 1

    print(f"final chunks: {len(chunks)} chunks + {len(uncovered_chunks)} uncovered_chunks + {len(table_chunks)} table_chunks")
    record_results_in_txt(chunks, docx_path, prefix=prefix_path, suffix='2chunks', split=True)
    record_results_in_txt(uncovered_chunks, docx_path, prefix=prefix_path, suffix='2chunks', mode='a+', split=True)
    record_results_in_txt(table_chunks, docx_path, prefix=prefix_path, suffix='2chunks', mode='a+', split=True)


# TODO fufill 前面的ID（如果跨行的话），然后按照ID分组（不知道是不是有序的），根据分组之后合并相同的内容，不相同的内容就列点；
#  同一个ID就构成一个语义块
def xlsx_process(file_path):
    prefix_path=f"xlsx_results_{file_path.split('/')[-1].split('.')[0]}"
    # 读取Excel文件
    df = pd.read_excel(file_path, sheet_name=None)
    # 遍历每个sheet
    for sheet_name, data in df.items():
        df_sheet = pd.DataFrame(data)
        results = []
        
        # 分组处理数据（合并id相同的内容）
        print(f"Processing sheet: {sheet_name}")
        df_sheet[['id']] = df_sheet[['id']].ffill()
        print(df_sheet.shape)
        grouped = df_sheet.groupby('id')
        
        # 遍历每个组
        for group_id, group_data in grouped:
            sam_row = {}
            to_merge = {}
            to_merge_len = []
            for column in group_data.columns:
                if column == "id":
                    continue
                unique_values = group_data[column].dropna().unique()
                if len(unique_values) == 1:
                    sam_row[column] = unique_values[0]
                elif len(unique_values) > 1:
                    to_merge[column] = group_data[column]
                    to_merge_len.append(len(to_merge[column]))
            
            # 处理描述一致的列内容
            same_results = []
            for k, v in sam_row.items():
                same_results.append(f"{k}: {v}")
            same_result_text = "\n".join(same_results)

            # 处理待合并的列内容
            if len(to_merge) == 0:
                results.append(same_result_text)
                continue
            elif len(set(to_merge_len)) == 1:
                to_merge_size = to_merge_len[0]
            else:
                print(f"Warning: {group_id} has inconsistent lengths in to_merge columns.")
                continue
                
            merged_results = []
            for j in range(to_merge_size):
                merged_row = []
                for key in list(to_merge.keys()):
                    merged_row.append(f"{key}: {to_merge[key].iloc[j]}")
                merged_results.append("\t ".join(merged_row))
            merged_result_text = "\n".join(merged_results)

            results.append(f"{same_result_text}\n{merged_result_text}")
        
        # 按文件名和sheet名保存结果
        chunks = []
        for i, r in enumerate(results):
            chunks.append({'segment_id': i, 'text': r})
        record_results_in_txt(chunks, file_path, prefix=prefix_path, suffix=f'_{sheet_name}2chunks', mode='w', split=True)


def main(file:str):
    # 示例用法
    # docx_path = "docx/NAIS map.docx"
    # file_path = "docx/OneMap_PRD_沿路搜索全程排序.docx"
    
    if file.endswith('.xlsx'):
        xlsx_process(file)
    if file.endswith('.docx'):
        docx_process(file)








