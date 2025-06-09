from docx import Document
import pdfplumber
from collections import defaultdict, Counter
from openai import OpenAI
from dotenv import load_dotenv
import os
import json
from docx.shared import Pt, Twips
import re
import pandas as pd

load_dotenv(os.path.expanduser("~/dot_env/openai.env"))


# TODO 自动生成的编号无法被docx库解析，需要额外处理
def infer_heading_level_by_font_and_indent(docx_path):
    doc = Document(docx_path)
    font_sizes = set()
    indent_levels = set()

    # 收集所有字体大小和缩进值
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size:
                font_sizes.add(run.font.size.pt)  # 获取字体大小（磅值）

        if paragraph.paragraph_format.left_indent:
            indent_levels.add(paragraph.paragraph_format.left_indent)  # 获取缩进（Twips）

    # 排序字体大小和缩进（从大到小）
    sorted_font_sizes = sorted(font_sizes, reverse=True)
    sorted_indents = sorted(indent_levels, reverse=True)

    # 推断层级
    structured_text = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        # 获取当前段落的字体大小（取第一个 run 的字体）
        current_font_size = None
        for run in paragraph.runs:
            if run.font.size:
                current_font_size = run.font.size.pt
                break

        # 获取当前段落的缩进
        current_indent = paragraph.paragraph_format.left_indent

        # 推断层级（基于字体大小）
        font_level = 0
        if current_font_size and sorted_font_sizes:
            font_level = sorted_font_sizes.index(current_font_size) + 1  # 1-based

        # 推断层级（基于缩进）
        indent_level = 0
        if current_indent and sorted_indents:
            indent_level = sorted_indents.index(current_indent) + 1  # 1-based

        # 综合层级（优先字体，其次缩进）
        final_level = max(font_level, indent_level)

        structured_text.append({
            "text": text,
            "font_size": current_font_size,
            "indent": current_indent,
            "level": final_level,
        })

    return structured_text


def docx_to_markdown(docx_path):
    doc = Document(docx_path)
    markdown_lines = []

    structured_text = infer_heading_level_by_font_and_indent(docx_path)
    max_level = max(item["level"] for item in structured_text)
    for para in structured_text:
        level = para['level']
        text = (re.sub(r"[\r\n\t]+"," ",para['text'])).strip()
        # 转换标题（Heading 1 → #, Heading 2 → ##）
        if level > 0 and level < max_level:
            markdown_lines.append(f"{'#' * level} {text}")
        # 普通段落
        else:
            markdown_lines.append(text)

    return markdown_lines, max_level


def docx_table_to_markdown(docx_path):
    doc = Document(docx_path)
    markdown_lines = []
    # 处理表格（转换为 Markdown 表格语法）
    for table in doc.tables:
        tb = []
        tb.append("\n| " + " | ".join(["Header"] * len(table.columns)) + " |")
        tb.append("| " + " | ".join(["---"] * len(table.columns)) + " |")
        for row in table.rows:
            tb.append("| " + " | ".join(cell.text.strip() for cell in row.cells) + " |")
        markdown_lines.append(tb)
    return markdown_lines

def pdf_to_markdown(pdf_path):
    """
    将 PDF 转为 Markdown 格式，修复空格和字符重复问题。
    返回 markdown_lines 和 max_level
    """
    markdown_lines = []
    max_level = 0

    with pdfplumber.open(pdf_path) as pdf:
        all_font_sizes = []
        
        # 先统计整篇 PDF 的所有字体大小
        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["size"])
            if words:
                all_font_sizes.extend([w["size"] for w in words if w["size"] is not None])

        if not all_font_sizes:
            print(f"警告: 在PDF '{pdf_path}' 中未找到任何有效的字体大小信息。")
            return [], 0

        # 找出字体大小的频率分布
        freq = Counter(all_font_sizes)
        distinct_sorted_sizes = sorted(freq.keys(), reverse=True)
        title_font_thresholds = distinct_sorted_sizes[:6]
        
        # 创建字体大小到标题级别的映射
        font_size_to_level = {}
        for i, size in enumerate(title_font_thresholds):
            font_size_to_level[size] = i + 1

        # 遍历每页内容
        for page in pdf.pages:
            words = page.extract_words(
                extra_attrs=["size", "fontname"],
                keep_blank_chars=False,
                use_text_flow=True  # 更好的处理文本流
            )
            
            if not words:
                continue
                
            # 按行分组（使用更精确的y坐标范围）
            line_groups = defaultdict(list)
            tolerance = 3  # 行高容差
            
            for word in words:
                # 计算单词的中间y坐标作为分组依据
                mid_y = (word["top"] + word["bottom"]) / 2
                # 找到最接近的y坐标组
                found_group = False
                for y in list(line_groups.keys()):
                    if abs(y - mid_y) <= tolerance:
                        line_groups[y].append(word)
                        found_group = True
                        break
                if not found_group:
                    line_groups[mid_y].append(word)
            
            # 按y坐标排序行
            sorted_lines = sorted(line_groups.items(), key=lambda x: x[0])
            
            for y, line_words in sorted_lines:
                # 按x坐标排序行内单词
                line_words.sort(key=lambda w: w["x0"])
                
                # 构建行文本（智能添加空格）
                line_text = ""
                prev_word = None
                
                for word in line_words:
                    text = word["text"]
                    
                    # 处理空格问题
                    if prev_word:
                        # 计算单词间距
                        space = word["x0"] - prev_word["x1"]
                        # 基于字体大小计算期望空格宽度
                        expected_space = prev_word["size"] * 0.3
                        
                        # 如果间距超过期望值，添加空格
                        if space > expected_space:
                            line_text += " "
                    
                    line_text += text
                    prev_word = word
                
                # 跳过空行
                if not line_text.strip():
                    continue
                
                # 标题检测逻辑
                font_sizes = [w["size"] for w in line_words if w["size"] is not None]
                if font_sizes:
                    max_font = max(font_sizes)
                    ends_with_punct = line_text.endswith(("：", "。", "；", ".", ";", ":", "？", "！", "?", "!"))
                    
                    # 检查是否为标题
                    if max_font in font_size_to_level and not ends_with_punct:
                        level = font_size_to_level[max_font]
                        max_level = max(max_level, level)
                        markdown_lines.append(f"{'#' * level} {line_text}")
                        continue
                
                # 普通文本
                markdown_lines.append(line_text)
    
    # 后处理：合并连续短行
    processed_lines = []
    i = 0
    while i < len(markdown_lines):
        line = markdown_lines[i]
        
        # 如果是标题，直接添加
        if line.startswith("#"):
            processed_lines.append(line)
            i += 1
            continue
        
        # 尝试合并后续短行
        merged_line = line
        j = i + 1
        while j < len(markdown_lines):
            next_line = markdown_lines[j]
            
            # 遇到标题或空行则停止合并
            if next_line.startswith("#") or not next_line.strip():
                break
                
            # 合并条件：当前行以标点结束或下一行很短
            if (line.endswith(("，", "。", ",", ".")) or 
                len(next_line) < 30 or 
                len(merged_line + " " + next_line) < 80):
                merged_line += " " + next_line
                j += 1
            else:
                break
        
        processed_lines.append(merged_line)
        i = j  # 跳过已合并的行
    
    return processed_lines, max_level

# 读取 txt 或 md
def txt_to_markdown(file_path):
    with open(file_path, encoding='utf-8') as f:
        return [line.strip() for line in f if line.strip()], 0

def markdown_to_markdown(file_path):
    lines = []
    max_level = 0
    
    with open(file_path, encoding='utf-8') as f:
        for line in f:
            # 保存原始行（保留行尾空格但去除换行符）
            clean_line = line.rstrip('\n')
            lines.append(clean_line)
            
            # 跳过空行
            if not clean_line.strip():
                continue
                
            # 检测标题行（以1-6个#开头，后跟空格）
            stripped = clean_line.lstrip()
            if stripped.startswith('#'):
                level = 0
                for char in stripped:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                # 验证标题格式（1-6个#后必须有空格）
                if 1 <= level <= 6 and level < len(stripped) and stripped[level].isspace():
                    max_level = max(max_level, level)
    
    return lines, max_level

def llm_segment(client, content: str, engine_type='qwen-max'):
    system_prompt = f"""
# 指示
## 任务介绍
- **目标**：你的任务是根据语义将文档内容分割成主题连贯的单元。
具有相同主题的连续文本片段应分组到相同的分割单元中，并在主题转换时创建新的分割单元。
注意，切割后的结果需覆盖所有的文本片段，不能遗漏任何片段。
- **数据**：输入数据是一系列由“\n\n”分隔的连续文本片段。每个连续文本片段前面都有一个编号N，格式为：[EX N]。
### 输出格式
- 以**JSONL（JSON行）**格式输出分割结果，每个分割单元的字符数不应超过1000。每个词典代表一个片段，由一个或多个关于同一主题的连续文本片段组成。
每个词典应包括以下关键字：
-**segment_id**：此段的索引，从0开始。
-**start_exchange_number**：此段中第一个连续文本片段的编号。
-**end_exchange_number**：此段中最后一个连续文本片段的编号。
-**num_exchange**：一个整数，表示此段中连续文本片段的数量，计算如下：end_exchange_number- start_exchange_number+1。

预期输出的示例格式，请勿输出JSON格式分割结果意外事件的任何描述信息：
{{"segment_id": 0，"start_exchange_number": 0，"end_exchange_number": 5，"num_exchange": 6}}
{{"segment_id": 1，"start_exchange_number": 6，"end_exchange_number": 8，"num_exchange": 3}}
"""
    system_item = {"role": "system", "content": system_prompt}
    messages = [system_item]
    messages += [{"role": "user", "content": content}]

    completion = client.chat.completions.create(
        model=engine_type,
        messages=messages,
        temperature=1.0,
        max_tokens=4096,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    if completion.choices[0].message.content:
        return completion.choices[0].message.content.strip()
    else:
        print('API Error!')
        return None


def document_level_segment(content_list, max_level, ex_flag=True):
    full_texts = []
    para = []
    prefix = [f"{'#' * l}" for l in range(1, max_level)]
    for i, c in enumerate(content_list):
        if ex_flag:
            item = f"[EX {i}] {c}"
        else:
            item = c
        c = c.strip()
        # 拆分段落，分段传入LLM
        para_flag = False
        for pre in prefix:
            if c.startswith(pre):
                para_flag = True
                break
        if para_flag:
            full_texts.append(para)
            para = [item]
        else:
            para.append(item)
    full_texts.append(para)
    return full_texts


def record_results_in_txt(results: list, file_path: str, prefix='', suffix='', mode='w'):
    # 创建新目录路径（确保路径正确）
    directory = os.path.dirname(os.path.abspath(__file__)) 
    new_dir = os.path.join(directory, prefix)
    
    # 如果目录不存在则创建
    os.makedirs(new_dir, exist_ok=True)
    
    # 创建新文件名（添加后缀并确保是.txt文件）
    filename = os.path.basename(file_path)
    
    new_filename = f"{filename.split('.')[0]}{suffix}.txt"
    
    # 完整的文件路径
    new_path = os.path.join(new_dir, new_filename)
    
    # 写入文件
    with open(new_path, mode, encoding='utf-8') as f:
        json.dump(results, f, ensure_ascii=False, indent=4)
    
    print(f"{len(results)} results have been saved to {new_path}.")
    return new_path


def find_uncovered_indices(covered_ranges, total_range_start=0, total_range_end=160):
    # Step 1: 合并重叠或连续的区间
    overlap = 0
    merged_ranges = []
    for current in sorted(covered_ranges, key=lambda x: x["start_exchange_number"]):
        if not merged_ranges:
            merged_ranges.append(current)
        else:
            last = merged_ranges[-1]
            # 如果当前区间与上一个区间重叠或连续，则合并
            if int(current["start_exchange_number"]) <= int(last["end_exchange_number"]) + 1:
                overlap += 1
                last["end_exchange_number"] = max(int(last["end_exchange_number"]), int(current["end_exchange_number"]))
            else:
                merged_ranges.append(current)

    # Step 2: 找出未被覆盖的区间
    uncovered = []
    prev_end = total_range_start - 1  # 初始化为起始点前一个数

    for interval in merged_ranges:
        start_no = int(interval["start_exchange_number"])
        end_no = int(interval["end_exchange_number"])

        # 如果当前区间的起点 > prev_end + 1，说明中间有未被覆盖的部分
        if start_no > prev_end + 1:
            uncovered.append((prev_end + 1, start_no - 1))
        prev_end = max(prev_end, end_no)

    # 检查最后一段是否覆盖到 total_range_end
    if prev_end < total_range_end:
        uncovered.append((prev_end + 1, total_range_end))

    # Step 3: 展开所有未被覆盖的索引（可选）
    uncovered_indices = []
    for (start_no, end_no) in uncovered:
        uncovered_indices.extend(range(start_no, end_no + 1))
    print(f"overlap: {overlap}")
    print(f"uncovered: {uncovered}")
    return uncovered


def deduplicate_and_sort_chunks(chunks, min_len=10):
    # 使用 set 存储 seen 键值对 (start, end)
    seen = set()
    deduplicated = []
    for chunk in chunks:
        if len(chunk['text']) <= min_len:
            continue
        key = (int(chunk["start_exchange_number"]), int(chunk["end_exchange_number"]))
        if key not in seen:
            seen.add(key)
            deduplicated.append(chunk)

    # 按 start_exchange_number 升序排序
    merged_chunks = merge_contained_chunks(deduplicated)
    sorted_chunks = sorted(merged_chunks, key=lambda x: (int(x["start_exchange_number"]), int(x['end_exchange_number'])))
    for i, chunk in enumerate(sorted_chunks):
        chunk["segment_id"] = i
    return sorted_chunks


def merge_contained_chunks(chunks, min_threshold=50, max_threshold=256):
    # 标记被合并的 chunk（即不再保留的）
    merged_indices = set()

    for i, chunk_i in enumerate(chunks):
        text_i = chunk_i["text"]
        if len(text_i) >= min_threshold:
            continue  # 只考虑短文本 chunk 被合并

        for j, chunk_j in enumerate(chunks):
            if i == j or j in merged_indices:
                continue
            text_j = chunk_j["text"]
            if text_i in text_j and len(text_j) <= max_threshold:
                # 找到包含它的母 chunk
                merged_indices.add(i)
                break  # 一个子 chunk 只合并一次

    # 返回未被合并的 chunk 列表
    print(f'remove {len(merged_indices)} chunks from {len(chunks)} chunks. {len(chunks) - len(merged_indices)} chunks remain.')
    return [chunk for idx, chunk in enumerate(chunks) if idx not in merged_indices]


def file_process(file_path, kb_name="results"):
    client = OpenAI(api_key=os.getenv("QWEN_API_KEY"), base_url=os.getenv("QWEN_API_BASE"))
    engine_type = "qwen-max"
    sufix = file_path.split('.')[-1]
    table_list = []
    if sufix=='docx' or sufix=='doc':
        content_list, max_level = docx_to_markdown(file_path) 
        table_list = docx_table_to_markdown(file_path)
    elif sufix == 'pdf':
        content_list, max_level = pdf_to_markdown(file_path)
    elif sufix == 'md':
        content_list, max_level = markdown_to_markdown(file_path)
    elif sufix == 'txt':
        content_list, max_level = txt_to_markdown(file_path)
    else :
        raise TypeError("只能解析.docx, .doc, .md, .txt文件")
    print(f"content_list: {len(content_list)}")
    print(f"table_list: {len(table_list)}")
    # record_results_in_txt(content_list, file_path, prefix=prefix_path, suffix='2markdown')
    # record_results_in_txt(table_list, file_path, prefix=prefix_path, suffix='2markdown', mode='a+')

    # 按文档层级进行粗粒度拆分，以便后续分块输入到LLM中
    coarse_chunk = []
    full_texts = document_level_segment(content_list, max_level)
    full_texts += table_list
    for p in full_texts:
        para_text = "\n\n".join(p)
        coarse_chunk.append({'num_exchange': len(p), 'text': para_text, 'len': len(para_text)})
    # record_results_in_txt(coarse_chunk, file_path, prefix=prefix_path, suffix='2coarse_chunk', split=True, enum=True)
    print(f"coarse segment results: {len(coarse_chunk)}")

    # 粗粒度拆分后，使用LLM进行细粒度拆分
    seg_ids = []
    for p in coarse_chunk:
        seg_json = llm_segment(client, p['text'], engine_type)
        seg_id = seg_json.strip().split("\n")
        seg_ids += seg_id
    print(f"LLM segment results: {len(seg_ids)}")
    # 读取LLM细粒度分块结果，得到有效的分块区间
    llm_seg_results = []
    for line in seg_ids:
        line = line.strip()
        if line and line[0] == '{' and line[-1] == '}':
            seg = json.loads(line)
            llm_seg_results.append(seg)
    print(f"filtered LLM segment results: {len(llm_seg_results)}")
    # record_results_in_txt(llm_seg_results, file_path, prefix=prefix_path, suffix='2ids')

    # 根据分块区间处理文档内容
    chunks = []
    chunk_idx = 0
    for seg in llm_seg_results:
        start = int(seg['start_exchange_number'])
        end = int(seg['end_exchange_number'])
        num_exchange = end - start + 1
        if num_exchange >= 40:
            print(f"warning range: {start} - {end}")
            long_content_list = content_list[start:end + 1]
            new_chunk_list = document_level_segment(long_content_list, max_level, ex_flag=False)
            nc_start = start
            for nc in new_chunk_list:
                nc_text = "\n\n".join(nc)
                nc_end = nc_start + len(nc) - 1
                chunks.append({"segment_id": chunk_idx, "num_exchange": len(nc),
                               "start_exchange_number": nc_start, "end_exchange_number": nc_end,  "text": nc_text})
                chunk_idx += 1
                nc_start = nc_end + 1
            continue
        merge_list = content_list[start:end + 1]
        merge_text = "\n".join(merge_list)
        chunks.append({"segment_id": chunk_idx, "num_exchange": num_exchange, "start_exchange_number": start,
                       "end_exchange_number": end,  "text": merge_text})
        chunk_idx += 1

    # 处理未覆盖的索引
    # TODO 后续或可根据 start_exchange_number来编号，这样子方便判断上下文
    uncovered_list = find_uncovered_indices(chunks.copy(), total_range_start=0, total_range_end=len(content_list) - 1)
    uncovered_chunks = []
    for start, end in uncovered_list:
        merge_list = content_list[start:end + 1]
        merge_text = "\n".join(merge_list)
        uncovered_chunks.append({"segment_id": chunk_idx, "num_exchange": end - start + 1,
                                 "start_exchange_number": start, "end_exchange_number": end,  "text": merge_text})
        chunk_idx += 1

    # 合并已有分块和未处理分块，去重并排序
    text_chunks = chunks + uncovered_chunks
    text_chunks = deduplicate_and_sort_chunks(text_chunks)
    find_uncovered_indices(text_chunks.copy(), total_range_start=0, total_range_end=len(content_list) - 1)

    # 处理表格
    # TODO 后续或可定位表格上下文，并据此重新编号
    table_chunks = []
    chunk_idx = len(text_chunks)
    for tb in table_list:
        tb_txt = "\n".join(tb)
        table_chunks.append({"segment_id": chunk_idx, "num_exchange": 1, "start_exchange_number": 0,
                             "end_exchange_number": 0,  "text": tb_txt})
        chunk_idx += 1

    print(f"chunks: {len(chunks)} + {len(uncovered_chunks)} + {len(table_chunks)}")
    print(f'final chunks: {len(text_chunks)} + {len(table_chunks)}')
    prefix_path=os.path.join("konwledgeBase",kb_name)
    text_chunks.extend(table_chunks)
    newPath = record_results_in_txt(text_chunks, file_path, prefix=prefix_path, suffix='2chunks')
    return newPath


        









